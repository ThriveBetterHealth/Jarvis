"""Document intelligence service."""

import json
import time
from typing import Optional
from uuid import UUID

from fastapi import UploadFile
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from models.document import AnalysisStatus, Document, DocumentAnalysis, DocumentType
from models.user import User
from services.ai.providers.anthropic_provider import AnthropicProvider
from services.ai.providers.base import ChatMessage
from services.file_service import FileService


MIME_TO_DOCTYPE = {
    "application/pdf": DocumentType.PDF,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentType.DOCX,
    "image/jpeg": DocumentType.IMAGE,
    "image/png": DocumentType.IMAGE,
    "image/tiff": DocumentType.IMAGE,
    "text/csv": DocumentType.CSV,
    "application/json": DocumentType.JSON,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": DocumentType.XLSX,
    "text/plain": DocumentType.TEXT,
}


class DocumentService:
    def __init__(self, db: AsyncSession):
        self.db = db

    async def list_documents(self, user_id: UUID, skip: int = 0, limit: int = 50) -> list:
        result = await self.db.execute(
            select(Document)
            .where(Document.user_id == user_id, Document.deleted_at.is_(None))
            .order_by(Document.created_at.desc())
            .offset(skip)
            .limit(limit)
        )
        return result.scalars().all()

    async def upload_and_analyse(self, user: User, file: UploadFile, analysis_type: str = "general") -> dict:
        file_service = FileService(self.db)
        stored_file = await file_service.upload(user=user, file=file)

        doc_type = MIME_TO_DOCTYPE.get(file.content_type, DocumentType.TEXT)
        doc = Document(
            user_id=user.id,
            file_id=stored_file.id,
            original_name=file.filename,
            document_type=doc_type,
            file_size_bytes=stored_file.size_bytes,
        )
        self.db.add(doc)
        await self.db.flush()

        # Extract text
        extracted_text = await self._extract_text(stored_file.storage_path, doc_type, file.filename)
        doc.extracted_text = extracted_text[:100000]  # Cap at 100k chars
        await self.db.flush()

        # Run analysis
        analysis = await self._run_analysis(doc, extracted_text, analysis_type)
        return {
            "document_id": str(doc.id),
            "file_id": str(stored_file.id),
            "analysis": {
                "summary": analysis.summary,
                "key_insights": analysis.key_insights,
                "risks": analysis.risks,
                "action_items": analysis.action_items,
            },
        }

    async def _extract_text(self, path: str, doc_type: DocumentType, filename: str) -> str:
        try:
            if doc_type == DocumentType.PDF:
                import PyPDF2
                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    return "\n".join(page.extract_text() or "" for page in reader.pages)
            elif doc_type == DocumentType.DOCX:
                from docx import Document as DocxDocument
                doc = DocxDocument(path)
                return "\n".join(p.text for p in doc.paragraphs)
            elif doc_type == DocumentType.IMAGE:
                import pytesseract
                from PIL import Image
                return pytesseract.image_to_string(Image.open(path))
            elif doc_type == DocumentType.CSV:
                import pandas as pd
                df = pd.read_csv(path)
                return df.to_string()
            elif doc_type == DocumentType.XLSX:
                import pandas as pd
                df = pd.read_excel(path)
                return df.to_string()
            else:
                with open(path, "r", errors="replace") as f:
                    return f.read()
        except Exception as e:
            return f"[Text extraction failed: {e}]"

    async def _run_analysis(self, doc: Document, text: str, analysis_type: str) -> DocumentAnalysis:
        analysis = DocumentAnalysis(
            document_id=doc.id,
            status=AnalysisStatus.PROCESSING,
            analysis_type=analysis_type,
        )
        self.db.add(analysis)
        await self.db.flush()

        start = time.time()
        try:
            provider = AnthropicProvider()
            prompt = f"""Analyse the following document and return a JSON object with:
- summary: one paragraph overview
- key_insights: array of top 5 findings
- risks: array of identified risks or concerns
- action_items: array of recommended actions

Return ONLY the JSON object.

Document content:
{text[:12000]}"""
            response = await provider.chat([
                ChatMessage(role="system", content="You are Jarvis, a document intelligence assistant."),
                ChatMessage(role="user", content=prompt),
            ])
            result = json.loads(response["content"])
            analysis.summary = result.get("summary", "")
            analysis.key_insights = result.get("key_insights", [])
            analysis.risks = result.get("risks", [])
            analysis.action_items = result.get("action_items", [])
            analysis.status = AnalysisStatus.COMPLETED
            analysis.model_used = "claude-sonnet-4-6"
        except Exception as e:
            analysis.status = AnalysisStatus.FAILED
            analysis.error_message = str(e)

        analysis.processing_time_ms = int((time.time() - start) * 1000)
        await self.db.flush()
        return analysis

    async def get_document(self, document_id: UUID, user_id: UUID) -> Optional[Document]:
        result = await self.db.execute(
            select(Document).where(
                Document.id == document_id,
                Document.user_id == user_id,
                Document.deleted_at.is_(None),
            )
        )
        return result.scalar_one_or_none()

    async def get_latest_analysis(self, document_id: UUID, user_id: UUID) -> Optional[DocumentAnalysis]:
        doc = await self.get_document(document_id, user_id)
        if not doc:
            return None
        result = await self.db.execute(
            select(DocumentAnalysis)
            .where(DocumentAnalysis.document_id == document_id)
            .order_by(DocumentAnalysis.created_at.desc())
            .limit(1)
        )
        return result.scalar_one_or_none()

    async def delete_document(self, document_id: UUID, user_id: UUID):
        doc = await self.get_document(document_id, user_id)
        if doc:
            doc.soft_delete()
            await self.db.flush()
